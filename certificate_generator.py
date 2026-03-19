import os
from docx import Document
from docx2pdf import convert
from docx.shared import RGBColor, Pt
from datetime import datetime
from server import app, send_certificates_email
from firebase_store import get_paper_by_id, update_paper

def replace_placeholders(doc, replacements):
    """Replace placeholders in all runs across paragraphs and tables with proper formatting."""
    
    # Define font settings
    font_name = "Times New Roman"
    font_size = Pt(16)
    
    def process_paragraph(paragraph):
        # First, combine all runs to ensure we can find split placeholders
        full_text = ''.join(run.text for run in paragraph.runs)
        
        # Check if any placeholder exists in this paragraph
        if not any(placeholder in full_text for placeholder in replacements.keys()):
            return
        
        # Clear all existing runs
        for run in paragraph.runs:
            run.text = ""
        
        # Process the text to find and format replacements
        current_index = 0
        formatted_runs = []
        
        # Find all placeholder positions and their formatting
        placeholder_positions = []
        for placeholder, (value, color) in replacements.items():
            start_idx = full_text.find(placeholder)
            while start_idx != -1:
                end_idx = start_idx + len(placeholder)
                placeholder_positions.append((start_idx, end_idx, value, color, placeholder))
                start_idx = full_text.find(placeholder, end_idx)
        
        # Sort by start index
        placeholder_positions.sort(key=lambda x: x[0])
        
        # Build the formatted text segments
        segments = []
        current_pos = 0
        
        for start, end, value, color, placeholder in placeholder_positions:
            # Add text before the placeholder
            if current_pos < start:
                segments.append((full_text[current_pos:start], None, False))
            
            # Add the replacement value with formatting
            segments.append((str(value), color, placeholder == "{{NAME}}"))
            current_pos = end
        
        # Add any remaining text after the last placeholder
        if current_pos < len(full_text):
            segments.append((full_text[current_pos:], None, False))
        
        # Create runs for each segment
        for text, color, is_bold in segments:
            if text:  # Only create run if there's text
                run = paragraph.add_run(text)
                run.font.name = font_name
                run.font.size = font_size
                if color:
                    run.font.color.rgb = color
                if is_bold:
                    run.bold = True

    # Process all paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    
    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

def generate_certificates(paper_id):
    with app.app_context():
        paper = get_paper_by_id(paper_id)
        if not paper:
            print(f"Paper ID {paper_id} not found!")
            return

        authors = [
            paper.get('corresponding_author_name'),
            paper.get('additional_author_name_1'),
            paper.get('additional_author_name_2'),
            paper.get('additional_author_name_3'),
            paper.get('additional_author_name_4'),
            paper.get('additional_author_name_5')
        ]
        authors = [name for name in authors if name]
        os.makedirs("certificates", exist_ok=True)
        pdf_paths = []

        # Format date properly
        if isinstance(paper.get('submission_date'), datetime):
            date_str = paper.get('submission_date').strftime("%d/%m/%y")
        else:
            date_str = str(paper.get('submission_date'))

        certificate_updates = {}

        for i, author_name in enumerate(authors):
            doc = Document("templates/certificate_template.docx")

            replacements = {
                "{{NAME}}": (author_name, RGBColor(0, 51, 153)),       # Blue color
                "{{TITLE}}": (paper.get('paper_title'), RGBColor(0, 51, 153)), # Blue color
                "{{ID}}": (paper.get('paper_id'), RGBColor(153, 0, 0)),        # Red color
                "{{DATE}}": (date_str, RGBColor(153, 0, 0))             # Red color
            }

            replace_placeholders(doc, replacements)

            docx_path = f"certificates/{paper.get('paper_id')}_author_{i+1}.docx"
            pdf_path = f"certificates/{paper.get('paper_id')}_author_{i+1}.pdf"
            doc.save(docx_path)

            try:
                convert(docx_path, pdf_path)
                os.remove(docx_path)  # remove word file if PDF created
            except Exception as e:
                print(f"docx2pdf failed: {e}. Keeping DOCX.")
                pdf_path = docx_path  # fallback if PDF fails

            pdf_paths.append(pdf_path)

            # Prepare cert path updates for Firestore
            if i == 0:
                certificate_updates['corresponding_author_certificate'] = pdf_path
            elif i == 1:
                certificate_updates['additional_author_cert_1'] = pdf_path
            elif i == 2:
                certificate_updates['additional_author_cert_2'] = pdf_path
            elif i == 3:
                certificate_updates['additional_author_cert_3'] = pdf_path
            elif i == 4:
                certificate_updates['additional_author_cert_4'] = pdf_path
            elif i == 5:
                certificate_updates['additional_author_cert_5'] = pdf_path

        if certificate_updates:
            update_paper(paper.get('paper_id'), certificate_updates)

        send_certificates_email(
            paper.get('corresponding_author_email'),
            paper.get('paper_id'),
            pdf_paths
        )

if __name__ == "__main__":
    paper_id = input("Enter Paper ID: ")
    generate_certificates(paper_id)